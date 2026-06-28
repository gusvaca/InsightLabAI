from docx import Document

from docx.shared import Pt

from analytics.insights.report_model import ExecutiveReport


class WordReport:

    """
    =======================================================
    Exportador Microsoft Word

    Convierte un ExecutiveReport en un documento DOCX.

    =======================================================
    """

    @staticmethod
    def generate(

        report: ExecutiveReport,

        output_file

    ):

        doc = Document()

        # ==========================================
        # TÍTULO
        # ==========================================

        titulo = doc.add_heading(

            report.titulo,

            level=1

        )

        titulo.style.font.size = Pt(24)

        if report.subtitulo:

            p = doc.add_paragraph()

            p.add_run(

                report.subtitulo

            ).italic = True

        # ==========================================
        # INFORMACIÓN GENERAL
        # ==========================================

        doc.add_heading(

            "Información General",

            level=2

        )

        tabla = doc.add_table(

            rows=4,

            cols=2

        )

        tabla.style = "Table Grid"

        tabla.cell(0,0).text = "Fecha"

        tabla.cell(0,1).text = report.fecha

        tabla.cell(1,0).text = "Score"

        tabla.cell(1,1).text = str(report.score)

        tabla.cell(2,0).text = "Estado"

        tabla.cell(2,1).text = report.estado

        tabla.cell(3,0).text = "Data Maturity"

        tabla.cell(3,1).text = report.data_maturity

        # ==========================================
        # SECCIONES
        # ==========================================

        for section in report.sections:

            doc.add_heading(

                section.titulo,

                level=2

            )

            doc.add_paragraph(

                section.contenido

            )

        # ==========================================
        # RECOMENDACIONES
        # ==========================================

        if report.recomendaciones:

            doc.add_heading(

                "Recomendaciones",

                level=2

            )

            for r in report.recomendaciones:

                doc.add_paragraph(

                    r,

                    style="List Bullet"

                )

        # ==========================================
        # SIGUIENTES PASOS
        # ==========================================

        if report.siguientes_pasos:

            doc.add_heading(

                "Próximos Pasos",

                level=2

            )

            for paso in report.siguientes_pasos:

                doc.add_paragraph(

                    paso,

                    style="List Number"

                )

        # ==========================================
        # PIE
        # ==========================================

        doc.add_page_break()

        doc.add_paragraph(

            "Reporte generado automáticamente por InsightLab AI Enterprise."

        )

        doc.save(

            output_file

        )

        return output_file